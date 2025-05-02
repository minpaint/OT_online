# directory/document_generators/siz_card_docx_generator.py
"""
🛡️ Генератор личной карточки учёта СИЗ (DOCX).

- Использует механизм маркеров аналогичный листу ознакомления
- Поддерживает объединенные строки для условий на лицевой стороне
- Заполняет только необходимые колонки на оборотной стороне
- Форматирует заголовки таблиц шрифтом 9 пунктов
- Применяет единый стиль форматирования для всех строк таблицы
- Устанавливает одинарный междустрочный интервал
- Убирает отступы до и после абзацев
- Генерирует случайные размеры СИЗ в зависимости от пола
"""

import logging
import traceback
import re
import random
from typing import Dict, Any, Optional, List, Tuple

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from directory.models.document_template import GeneratedDocument
from directory.document_generators.base import (
    get_document_template,
    prepare_employee_context,
    generate_docx_from_template,
)
from directory.models.siz import SIZNorm
from directory.models.siz_issued import SIZIssued

logger = logging.getLogger(__name__)


# =============================================
# Основная функция генерации карточки СИЗ
# =============================================

def generate_siz_card_docx(
        employee,
        user=None,
        custom_context: Optional[Dict[str, Any]] = None,
) -> Optional[GeneratedDocument]:
    """Генерирует карточку учёта СИЗ для сотрудника."""

    try:
        # 1. Получение шаблона
        template = get_document_template("siz_card", employee)
        if not template:
            raise ValueError("Активный шаблон для карточки СИЗ не найден")

        # 2. Подготовка базового контекста
        context = prepare_employee_context(employee)
        full_name = context.get("fio_nominative", "")

        # Разделяем полное имя на составляющие
        last_name, first_name, patronymic = _split_full_name(full_name)

        # 3. Определение пола для заголовка
        gender = _gender_from_patronymic(patronymic)

        # 4. Генерация случайных размеров СИЗ
        ppe_head, ppe_gloves, sizod, _ = _generate_random_ppe_sizes(gender)

        # 5. Получение выбранных норм СИЗ
        selected_norm_ids = []
        if custom_context:
            if 'selected_norm_ids' in custom_context:
                selected_norm_ids = custom_context['selected_norm_ids']
            elif 'selected_norms' in custom_context:
                selected_norm_ids = custom_context['selected_norms']

        # 6. Если нормы не выбраны, включаем все нормы для должности
        if not selected_norm_ids and hasattr(employee, 'position') and employee.position:
            logger.info("Нормы не выбраны, используем все нормы для должности")
            selected_norm_ids = list(SIZNorm.objects.filter(
                position=employee.position
            ).values_list('id', flat=True))

        # 7. Получение данных норм СИЗ
        norms_data = []
        if employee.position and selected_norm_ids:
            norm_query = SIZNorm.objects.filter(
                id__in=selected_norm_ids
            ).select_related('siz')

            for norm in norm_query:
                norms_data.append({
                    "name": norm.siz.name,
                    "classification": norm.siz.classification,
                    "unit": norm.siz.unit,
                    "quantity": norm.quantity,
                    "wear_period": "До износа" if norm.siz.wear_period == 0 else str(norm.siz.wear_period),
                    "condition": norm.condition  # Добавляем условие для группировки
                })

            logger.info(f"Найдено {len(norms_data)} норм СИЗ для сотрудника")

        # 8. Формирование контекста с заглушками для недостающих данных
        department_name = context.get("department", "")
        subdivision_name = context.get("subdivision", "")
        organization_name = context.get("organization_name", "")

        # Очищаем подразделение от дополнительной информации, если она есть
        if subdivision_name and "(" in subdivision_name:
            # Оставляем только название до скобки
            subdivision_name = subdivision_name.split("(")[0].strip()

        context.update({
            "card_number": f"SIZ-{employee.id}",
            "employee_full_name": full_name,
            "last_name": last_name,
            "first_name": first_name,
            "patronymic": patronymic,
            "employee_gender": gender,
            "employee_height": getattr(employee, "height", "") or "",
            "employee_clothing_size": getattr(employee, "clothing_size", "") or "",
            "employee_shoe_size": getattr(employee, "shoe_size", "") or "",
            "department_name": department_name,
            "subdivision_name": subdivision_name,  # Добавляем подразделение в контекст
            "organization_name": organization_name,  # Добавляем организацию в контекст
            "position_name": context.get("position_nominative", ""),
            "hire_date": employee.hire_date.strftime("%d.%m.%Y") if hasattr(employee,
                                                                            "hire_date") and employee.hire_date else "",
            # Размеры СИЗ
            "ppe_head": ppe_head,
            "ppe_gloves": ppe_gloves,
            "sizod": sizod,

            # Маркеры для таблиц
            "NORMS_TABLE": "NORMS_TABLE_MARKER",
            "ISSUED_TABLE": "ISSUED_TABLE_MARKER",

            # Данные для таблиц
            "siz_norms": norms_data,
            "issued_siz": norms_data  # Используем те же данные для оборотной стороны
        })

        # 9. Добавление пользовательского контекста
        if custom_context:
            for k, v in custom_context.items():
                if k not in ['selected_norm_ids', 'selected_norms']:
                    context[k] = v

        # 10. Генерация документа + пост-обработка
        return generate_docx_from_template(
            template,
            context,
            employee,
            user,
            post_processor=process_siz_card_tables,
        )

    except Exception as exc:
        logger.error("Ошибка при генерации карточки СИЗ: %s", exc)
        logger.error(traceback.format_exc())
        return None


# =========================
#   ГЕНЕРАЦИЯ РАЗМЕРОВ СИЗ
# =========================

def _split_full_name(full_name: str) -> Tuple[str, str, str]:
    """Разделяет полное имя на фамилию, имя и отчество."""
    parts = full_name.split()

    if len(parts) >= 3:
        return parts[0], parts[1], parts[2]
    elif len(parts) == 2:
        return parts[0], parts[1], ""
    elif len(parts) == 1:
        return parts[0], "", ""
    else:
        return "", "", ""


def _generate_random_ppe_sizes(gender: str) -> Tuple[str, str, str, str]:
    """
    Генерирует случайные размеры СИЗ в зависимости от пола.

    Args:
        gender: Пол сотрудника ("Мужской" или "Женский")

    Returns:
        Кортеж (headgear, gloves, respirator, gas_mask)
    """
    if gender == "Мужской":
        # Мужские размеры
        headgear = random.randint(55, 59)  # Головной убор от 55 до 59
        gloves = random.randint(15, 19) / 2  # Перчатки от 7.5 до 9.5, кратные 0.5
        respirator = random.choice(["1", "2", "3"])  # Респиратор размеры 1, 2, 3
    else:
        # Женские размеры
        headgear = random.randint(53, 57)  # Головной убор от 53 до 57
        gloves = random.randint(13, 17) / 2  # Перчатки от 6.5 до 8.5, кратные 0.5
        respirator = random.choice(["1", "2", "3"])  # Респиратор размеры 1, 2, 3

    # Противогаз такого же размера, как и респиратор
    gas_mask = respirator

    return str(headgear), str(gloves), respirator, gas_mask


# =========================
#   ПОСТ-ОБРАБОТКА ТАБЛИЦ
# =========================

def process_siz_card_tables(doc, context):
    """Обрабатывает таблицы в сгенерированном документе."""
    try:
        # Получаем данные норм СИЗ
        norms_data = context.get("siz_norms", [])
        if not norms_data:
            logger.warning("Нет данных о нормах СИЗ для обработки таблиц")
            return doc

        docx_document = doc.docx

        # Обрабатываем лицевую сторону
        processed_front = False
        for table in docx_document.tables:
            row_idx, cell_idx = _find_marker_in_table(table, "NORMS_TABLE_MARKER")
            if row_idx is not None:
                # Обработка таблицы лицевой стороны
                process_front_table(table, row_idx, cell_idx, norms_data)
                processed_front = True
                break

        if not processed_front:
            logger.warning("Маркер NORMS_TABLE_MARKER не найден в таблицах")

        # Обрабатываем оборотную сторону
        processed_back = False
        for table in docx_document.tables:
            row_idx, cell_idx = _find_marker_in_table(table, "ISSUED_TABLE_MARKER")
            if row_idx is not None:
                # Обработка таблицы оборотной стороны
                process_back_table(table, row_idx, cell_idx, norms_data)
                processed_back = True
                break

        if not processed_back:
            logger.warning("Маркер ISSUED_TABLE_MARKER не найден в таблицах")

        return doc

    except Exception as exc:
        logger.error("Ошибка при пост-обработке таблиц: %s", exc)
        logger.error(traceback.format_exc())
        return doc


def process_front_table(table, row_idx, cell_idx, norms_data):
    """Обрабатывает таблицу на лицевой стороне с объединёнными строками для условий."""
    try:
        # Группируем нормы по условиям
        grouped_norms = {}
        for norm in norms_data:
            condition = norm.get("condition", "")
            if condition not in grouped_norms:
                grouped_norms[condition] = []
            grouped_norms[condition].append(norm)

        # Удаляем маркер из ячейки
        cell = table.rows[row_idx].cells[cell_idx]
        cell.text = ""

        # Форматируем заголовки таблицы - уменьшаем размер шрифта до 9 пт
        if row_idx > 0:  # Предполагаем, что заголовок находится в строке выше маркера
            header_row = table.rows[row_idx - 1]
            _format_header_row(header_row)

        # Заполняем первую строку, если есть нормы без условий
        current_row = row_idx
        template_row = None
        if current_row < len(table.rows):
            template_row = table.rows[current_row]

        if "" in grouped_norms and grouped_norms[""]:
            for norm in grouped_norms[""]:
                if current_row >= len(table.rows):
                    new_row = table.add_row()
                    # Копируем формат с образца, если он есть
                    if template_row:
                        _copy_row_properties(template_row, new_row)
                else:
                    new_row = table.rows[current_row]

                _fill_front_row(new_row, norm)
                current_row += 1

        # Добавляем строки с условиями и соответствующими нормами
        for condition, norms in grouped_norms.items():
            if not condition:
                continue  # Пропускаем пустые условия, они уже обработаны

            # Добавляем строку с условием
            if current_row >= len(table.rows):
                condition_row = table.add_row()
                # Если есть шаблонная строка, копируем её свойства
                if template_row:
                    _copy_row_properties(template_row, condition_row)
            else:
                condition_row = table.rows[current_row]

            # Объединяем ячейки в строке условия
            first_cell = condition_row.cells[0]
            for i in range(1, len(condition_row.cells)):
                if i < len(condition_row.cells):
                    try:
                        first_cell.merge(condition_row.cells[i])
                    except Exception as e:
                        logger.warning(f"Не удалось объединить ячейки: {str(e)}")

            # Заполняем текст условия
            first_cell.text = condition
            for paragraph in first_cell.paragraphs:
                # Устанавливаем междустрочный интервал как одинарный
                paragraph.paragraph_format.line_spacing = 1.0
                # Убираем отступы до и после абзаца
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run.italic = True

            current_row += 1

            # Добавляем нормы для данного условия
            for norm in norms:
                if current_row >= len(table.rows):
                    new_row = table.add_row()
                    # Копируем формат с образца, если он есть
                    if template_row:
                        _copy_row_properties(template_row, new_row)
                else:
                    new_row = table.rows[current_row]

                _fill_front_row(new_row, norm)
                current_row += 1

        # Применяем форматирование ко всей таблице
        _apply_table_format(table)

    except Exception as e:
        logger.error(f"Ошибка при обработке таблицы лицевой стороны: {str(e)}")
        logger.error(traceback.format_exc())


def process_back_table(table, row_idx, cell_idx, norms_data):
    """Обрабатывает таблицу на оборотной стороне, заполняя только нужные колонки."""
    try:
        # Удаляем маркер из ячейки
        cell = table.rows[row_idx].cells[cell_idx]
        cell.text = ""

        # Форматируем заголовки таблицы - уменьшаем размер шрифта до 9 пт
        if row_idx > 0:  # Предполагаем, что заголовок находится в строке выше маркера
            header_row = table.rows[row_idx - 1]
            _format_header_row(header_row)

            # Если есть строка с подзаголовками, форматируем и её
            if row_idx > 1:
                subheader_row = table.rows[row_idx - 2]
                _format_header_row(subheader_row)

        # Определяем колонки для заполнения (0-based): 0, 1, 3, 6
        cols_to_fill = [0, 1, 3, 6]

        # Создаем образец строки с правильным форматированием
        template_row = None
        if row_idx < len(table.rows):
            template_row = table.rows[row_idx]

        # Заполняем строки данными
        current_row = row_idx
        for norm in norms_data:
            if current_row >= len(table.rows):
                new_row = table.add_row()
                # Копируем формат с образца, если он есть
                if template_row:
                    _copy_row_properties(template_row, new_row)
            else:
                new_row = table.rows[current_row]

            # Заполняем только нужные колонки
            if 0 < len(new_row.cells):
                new_row.cells[0].text = norm.get("name", "")
                # Выравнивание по левому краю для первой ячейки
                for paragraph in new_row.cells[0].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 1 < len(new_row.cells):
                new_row.cells[1].text = norm.get("classification", "")
                # Центрирование для остальных ячеек
                for paragraph in new_row.cells[1].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 3 < len(new_row.cells):
                new_row.cells[3].text = str(norm.get("quantity", ""))
                for paragraph in new_row.cells[3].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 6 < len(new_row.cells):
                new_row.cells[6].text = "✓"  # Галочка в колонке расписки
                for paragraph in new_row.cells[6].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            # Форматируем ячейки
            for col in cols_to_fill:
                if col < len(new_row.cells):
                    _set_cell_border(new_row.cells[col])

            current_row += 1

        # Применяем форматирование ко всей таблице
        _apply_table_format(table)

    except Exception as e:
        logger.error(f"Ошибка при обработке таблицы оборотной стороны: {str(e)}")
        logger.error(traceback.format_exc())


# --------------------------
#   УТИЛИТЫ
# --------------------------

def _copy_row_properties(src_row, dst_row):
    """Копирует все свойства из исходной строки в целевую."""
    try:
        # Копируем высоту строки и другие атрибуты
        dst_row.height = src_row.height

        # Копируем количество ячеек и их свойства
        for i in range(min(len(src_row.cells), len(dst_row.cells))):
            # Сбрасываем текст, но копируем формат
            dst_row.cells[i].text = ""

            # Пытаемся копировать стиль ячейки и её границы
            try:
                dst_tc = dst_row.cells[i]._tc
                src_tc = src_row.cells[i]._tc

                # Копируем все свойства таблицы
                if src_tc.tcPr is not None:
                    if dst_tc.tcPr is None:
                        dst_tc.tcPr = OxmlElement('w:tcPr')

                    # Копируем ширину ячейки
                    src_width = src_tc.tcPr.find(qn('w:tcW'))
                    if src_width is not None:
                        dst_width = dst_tc.tcPr.find(qn('w:tcW'))
                        if dst_width is None:
                            dst_width = OxmlElement('w:tcW')
                            dst_tc.tcPr.append(dst_width)
                        dst_width.attrib.update(src_width.attrib)
            except Exception as e:
                logger.debug(f"Не удалось скопировать свойства ячейки: {str(e)}")
                # Продолжаем выполнение

            # Устанавливаем границы
            _set_cell_border(dst_row.cells[i])
    except Exception as e:
        logger.error(f"Ошибка при копировании свойств строки: {str(e)}")


def _format_header_row(header_row):
    """Форматирует строку заголовка таблицы."""
    try:
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                # Устанавливаем междустрочный интервал как одинарный
                paragraph.paragraph_format.line_spacing = 1.0
                # Убираем отступы до и после абзаца
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)  # Уменьшаем размер до 9 пунктов
                    run.font.bold = True  # Делаем заголовки полужирными
    except Exception as e:
        logger.error(f"Ошибка при форматировании заголовков: {str(e)}")


def _gender_from_patronymic(patronymic: str) -> str:
    """Определяет пол по отчеству."""
    if not patronymic:
        return "Мужской"  # По умолчанию

    if patronymic.endswith(("на", "вна", "чна", "кызы", "зы")):
        return "Женский"
    if patronymic.endswith(("ич", "ыч", "оглы", "улы", "лы")):
        return "Мужской"

    # По умолчанию считаем мужским
    return "Мужской"


def _find_marker_in_table(table, marker):
    """Находит маркер в таблице и возвращает координаты ячейки."""
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if marker in paragraph.text:
                    return r_idx, c_idx
    return None, None


def _fill_front_row(row, norm):
    """Заполняет строку таблицы на лицевой стороне."""
    # Проверяем наличие достаточного количества ячеек
    if len(row.cells) < 5:
        logger.warning(f"Недостаточно ячеек в строке таблицы: {len(row.cells)}")
        return

    # Очищаем ячейки перед заполнением для предотвращения проблем с форматированием
    for i in range(5):
        row.cells[i].text = ""

    # Заполняем ячейки
    row.cells[0].text = norm.get("name", "")
    row.cells[1].text = norm.get("classification", "")
    row.cells[2].text = norm.get("unit", "")
    row.cells[3].text = str(norm.get("quantity", ""))
    row.cells[4].text = norm.get("wear_period", "")

    # Форматируем ячейки
    for i in range(5):
        for paragraph in row.cells[i].paragraphs:
            # Устанавливаем междустрочный интервал как одинарный
            paragraph.paragraph_format.line_spacing = 1.0
            # Убираем отступы до и после абзаца
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            # Первый столбец с наименованием СИЗ выравниваем по левому краю
            if i == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Остальные столбцы - по центру
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Устанавливаем шрифт для всего содержимого ячейки
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # Применяем границы
        _set_cell_border(row.cells[i])


def _apply_table_format(table):
    """Применяет форматирование к таблице."""
    # 1. Попытка применить готовый стиль
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        logger.info("Стиль 'Table Grid' недоступен – проставляем границы вручную")
        _add_borders_manually(table)

    # 2. Шрифт всем run‑ам, кроме заголовков, которые форматируются отдельно
    for row_idx, row in enumerate(table.rows):
        # Пропускаем заголовки, которые форматируются отдельно
        if row_idx < 2:  # Предполагаем, что заголовки в первых двух строках
            continue

        for cell in row.cells:
            for para in cell.paragraphs:
                # Устанавливаем междустрочный интервал как одинарный
                para.paragraph_format.line_spacing = 1.0
                # Убираем отступы до и после абзаца
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                if not para.runs:
                    para.add_run("")
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def _add_borders_manually(table):
    """Добавляет границы всем ячейкам таблицы."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell)


def _set_cell_border(
        cell,
        **kwargs,
):
    """Устанавливает одинарную границу толщиной 4 εм (≈0.5 pt) вокруг ячейки."""
    # Значения по умолчанию – все стороны single 4 εм
    sides = {
        "top": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "left": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "bottom": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "right": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
    }
    sides.update(kwargs)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tblBorders = tcPr.find(qn("w:tcBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tcBorders")
        tcPr.append(tblBorders)

    for edge in ("left", "top", "right", "bottom"):
        edge_el = tblBorders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tblBorders.append(edge_el)
        attrs = sides.get(edge, {})
        for key in ["val", "sz", "color", "space"]:
            edge_el.set(qn(f"w:{key}"), attrs.get(key, "single" if key == "val" else "4"))