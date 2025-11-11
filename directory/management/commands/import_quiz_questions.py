# directory/management/commands/import_quiz_questions.py
"""
Django management команда для импорта вопросов экзамена из Word документа.

Использование:
    python manage.py import_quiz_questions <путь_к_файлу.docx> --category="Название раздела"

Опции:
    --category: Название раздела для импорта (обязательно)
    --correct-answer: Номер правильного ответа (1-4) для всех вопросов (опционально)
    --dry-run: Тестовый запуск без сохранения в БД
    --clear: Очистить существующие вопросы в разделе перед импортом
"""

from django.core.management.base import BaseCommand, CommandError
from docx import Document
import re
from directory.models import QuizCategory, Question, Answer


class Command(BaseCommand):
    help = 'Импорт вопросов экзамена из Word документа'

    def add_arguments(self, parser):
        parser.add_argument(
            'docx_file',
            type=str,
            help='Путь к Word файлу с вопросами'
        )
        parser.add_argument(
            '--category',
            type=str,
            required=True,
            help='Название раздела для импорта вопросов'
        )
        parser.add_argument(
            '--correct-answer',
            type=int,
            choices=[1, 2, 3, 4],
            help='Номер правильного ответа (1-4) для всех вопросов'
        )
        parser.add_argument(
            '--dry-run',
            action='store_true',
            help='Тестовый запуск без сохранения в БД'
        )
        parser.add_argument(
            '--clear',
            action='store_true',
            help='Очистить существующие вопросы в разделе'
        )

    def handle(self, *args, **options):
        docx_file = options['docx_file']
        category_name = options['category']
        correct_answer_num = options.get('correct_answer')
        dry_run = options['dry_run']
        clear = options['clear']

        # Настройка кодировки для вывода
        import sys
        if sys.platform == 'win32':
            import io
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

        self.stdout.write(self.style.SUCCESS(f'\n{"="*60}'))
        self.stdout.write(self.style.SUCCESS('Импорт вопросов из Word документа'))
        self.stdout.write(self.style.SUCCESS(f'{"="*60}\n'))

        # Получаем или создаем раздел
        if not dry_run:
            category, created = QuizCategory.objects.get_or_create(
                name=category_name,
                defaults={'description': ''}
            )
            if created:
                self.stdout.write(self.style.SUCCESS(f'[+] Создан новый раздел: {category_name}'))
            else:
                self.stdout.write(self.style.WARNING(f'[*] Используется существующий раздел: {category_name}'))

            # Очистка существующих вопросов
            if clear:
                deleted_count = Question.objects.filter(category=category).delete()[0]
                self.stdout.write(self.style.WARNING(f'[-] Удалено {deleted_count} существующих вопросов'))
        else:
            self.stdout.write(self.style.WARNING('ТЕСТОВЫЙ РЕЖИМ (dry-run) - данные не будут сохранены'))
            category = None

        # Загружаем документ
        try:
            doc = Document(docx_file)
            self.stdout.write(self.style.SUCCESS(f'\n[OK] Документ загружен: {docx_file}'))
            self.stdout.write(f'  Параграфов: {len(doc.paragraphs)}\n')
        except Exception as e:
            raise CommandError(f'Ошибка при открытии файла: {e}')

        # Парсим вопросы
        questions = self.parse_questions(doc)
        self.stdout.write(self.style.SUCCESS(f'[OK] Распознано вопросов: {len(questions)}\n'))

        # Импортируем вопросы
        imported_count = 0
        errors = []

        for idx, question_data in enumerate(questions, 1):
            try:
                if dry_run:
                    # Просто выводим информацию
                    self.stdout.write(f'\n[{idx}] {question_data["text"][:80]}...')
                    self.stdout.write(f'    Вариантов ответов: {len(question_data["answers"])}')
                    for i, answer in enumerate(question_data["answers"], 1):
                        marker = ' [OK]' if correct_answer_num == i else ''
                        self.stdout.write(f'      {i}. {answer[:60]}{marker}')
                else:
                    # Создаем вопрос
                    question = Question.objects.create(
                        category=category,
                        question_text=question_data['text'],
                        order=idx
                    )

                    # Создаем варианты ответов
                    for i, answer_text in enumerate(question_data['answers'], 1):
                        is_correct = (correct_answer_num == i) if correct_answer_num else False
                        Answer.objects.create(
                            question=question,
                            answer_text=answer_text,
                            is_correct=is_correct,
                            order=i
                        )

                    imported_count += 1

                    if idx % 10 == 0:
                        self.stdout.write(self.style.SUCCESS(f'[OK] Импортировано {idx} вопросов...'))

            except Exception as e:
                error_msg = f'Ошибка при импорте вопроса #{idx}: {e}'
                errors.append(error_msg)
                self.stdout.write(self.style.ERROR(f'[ERROR] {error_msg}'))

        # Итоги
        self.stdout.write(self.style.SUCCESS(f'\n{"="*60}'))
        if dry_run:
            self.stdout.write(self.style.WARNING('ТЕСТОВЫЙ РЕЖИМ - данные не сохранены'))
            self.stdout.write(f'Всего вопросов для импорта: {len(questions)}')
        else:
            self.stdout.write(self.style.SUCCESS(f'[OK] Успешно импортировано вопросов: {imported_count}'))
            if errors:
                self.stdout.write(self.style.ERROR(f'[ERROR] Ошибок: {len(errors)}'))
                for error in errors[:5]:  # Показываем первые 5 ошибок
                    self.stdout.write(self.style.ERROR(f'  - {error}'))
        self.stdout.write(self.style.SUCCESS(f'{"="*60}\n'))

    def parse_questions(self, doc):
        """
        Парсит вопросы из Word документа.

        Ожидаемый формат:
        - Параграф с "Вопрос № X из Y" и текстом вопроса
        - Следующие параграфы - варианты ответов (до следующего вопроса)
        """
        questions = []
        current_question = None

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Проверяем, начинается ли новый вопрос
            if 'Вопрос №' in text and 'из' in text:
                # Сохраняем предыдущий вопрос
                if current_question and current_question['answers']:
                    questions.append(current_question)

                # Извлекаем текст вопроса
                # Формат: "Вопрос № X из Y\nВопрос: текст вопроса\nВарианты ответа:"
                lines = text.split('\n')
                question_text = ''

                for line in lines:
                    if line.strip().startswith('Вопрос:'):
                        # Убираем "Вопрос:" из начала
                        question_text = line.replace('Вопрос:', '').strip()
                        break

                # Если не нашли "Вопрос:", берем всё после номера
                if not question_text:
                    # Находим "Вопрос № X из Y" и берем всё после него
                    match = re.search(r'Вопрос № \d+ из \d+\s*(.*)', text, re.DOTALL)
                    if match:
                        full_text = match.group(1).strip()
                        # Убираем "Варианты ответа:" если есть
                        full_text = re.sub(r'Варианты ответа:.*$', '', full_text, flags=re.DOTALL).strip()
                        question_text = full_text

                current_question = {
                    'text': question_text,
                    'answers': []
                }

            elif current_question is not None:
                # Это вариант ответа
                # Пропускаем строки с "Варианты ответа:"
                if 'Варианты ответа:' not in text and 'Вопрос:' not in text:
                    current_question['answers'].append(text)

        # Добавляем последний вопрос
        if current_question and current_question['answers']:
            questions.append(current_question)

        return questions
