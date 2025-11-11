# directory/management/commands/import_quiz_questions_v2.py
"""
Улучшенная Django management команда для импорта вопросов из Word документа.

НОВЫЕ ВОЗМОЖНОСТИ:
- Автоматическое распознавание правильного ответа по жирному шрифту
- Извлечение источника (пояснения) к вопросу
- Поддержка нескольких правильных ответов
- Детальная валидация структуры

Использование:
    python manage.py import_quiz_questions_v2 <путь_к_файлу.docx> --category="Название раздела"

Опции:
    --category: Название раздела для импорта (обязательно)
    --dry-run: Тестовый запуск без сохранения в БД
    --clear: Очистить существующие вопросы в разделе перед импортом
    --verbose: Подробный вывод процесса парсинга
"""

from django.core.management.base import BaseCommand, CommandError
from docx import Document
import re
from directory.models import QuizCategory, Question, Answer


class Command(BaseCommand):
    help = 'Импорт вопросов экзамена из Word документа (v2 с автоопределением правильных ответов)'

    def add_arguments(self, parser):
        parser.add_argument(
            'docx_file',
            type=str,
            help='Путь к Word файлу с вопросами'
        )
        parser.add_argument(
            '--category',
            type=str,
            required=True,
            help='Название раздела для импорта вопросов'
        )
        parser.add_argument(
            '--dry-run',
            action='store_true',
            help='Тестовый запуск без сохранения в БД'
        )
        parser.add_argument(
            '--clear',
            action='store_true',
            help='Очистить существующие вопросы в разделе'
        )
        parser.add_argument(
            '--verbose',
            action='store_true',
            help='Подробный вывод'
        )

    def handle(self, *args, **options):
        docx_file = options['docx_file']
        category_name = options['category']
        dry_run = options['dry_run']
        clear = options['clear']
        self.verbose = options['verbose']

        # Настройка кодировки для вывода
        import sys
        if sys.platform == 'win32':
            import io
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

        self.stdout.write(self.style.SUCCESS(f'\n{"="*70}'))
        self.stdout.write(self.style.SUCCESS('Импорт вопросов из Word (v2 - с автоопределением)'))
        self.stdout.write(self.style.SUCCESS(f'{"="*70}\n'))

        # Получаем или создаем раздел
        if not dry_run:
            category, created = QuizCategory.objects.get_or_create(
                name=category_name,
                defaults={'description': ''}
            )
            if created:
                self.stdout.write(self.style.SUCCESS(f'[+] Создан новый раздел: {category_name}'))
            else:
                self.stdout.write(self.style.WARNING(f'[*] Используется существующий раздел: {category_name}'))

            # Очистка существующих вопросов
            if clear:
                deleted_count = Question.objects.filter(category=category).delete()[0]
                self.stdout.write(self.style.WARNING(f'[-] Удалено {deleted_count} существующих вопросов'))
        else:
            self.stdout.write(self.style.WARNING('ТЕСТОВЫЙ РЕЖИМ (dry-run) - данные не будут сохранены'))
            category = None

        # Загружаем документ
        try:
            doc = Document(docx_file)
            self.stdout.write(self.style.SUCCESS(f'\n[OK] Документ загружен: {docx_file}'))
            self.stdout.write(f'  Параграфов: {len(doc.paragraphs)}\n')
        except Exception as e:
            raise CommandError(f'Ошибка при открытии файла: {e}')

        # Парсим вопросы
        questions = self.parse_questions_v2(doc)
        self.stdout.write(self.style.SUCCESS(f'[OK] Распознано вопросов: {len(questions)}\n'))

        # Статистика
        questions_with_sources = sum(1 for q in questions if q.get('explanation'))
        questions_with_correct = sum(1 for q in questions if any(a['is_correct'] for a in q['answers']))

        self.stdout.write(f'  • С источниками (пояснениями): {questions_with_sources}')
        self.stdout.write(f'  • С отмеченными правильными ответами: {questions_with_correct}\n')

        # Импортируем вопросы
        imported_count = 0
        warnings = []
        errors = []

        for idx, question_data in enumerate(questions, 1):
            try:
                # Проверка: есть ли правильный ответ?
                has_correct = any(a['is_correct'] for a in question_data['answers'])
                if not has_correct:
                    warnings.append(f'Вопрос #{idx}: нет правильного ответа (жирного текста)')

                if dry_run:
                    # Просто выводим информацию
                    self.stdout.write(f'\n[{idx}] {question_data["text"][:80]}...')
                    self.stdout.write(f'    Вариантов ответов: {len(question_data["answers"])}')

                    for i, answer in enumerate(question_data["answers"], 1):
                        marker = ' ✓ ПРАВИЛЬНЫЙ' if answer['is_correct'] else ''
                        bold_marker = ' [BOLD]' if answer['is_bold'] else ''
                        self.stdout.write(f'      {i}. {answer["text"][:60]}{marker}{bold_marker}')

                    if question_data.get('explanation'):
                        self.stdout.write(f'    📚 Источник: {question_data["explanation"][:80]}...')
                else:
                    # Создаем вопрос
                    question = Question.objects.create(
                        category=category,
                        question_text=question_data['text'],
                        explanation=question_data.get('explanation', ''),
                        order=idx
                    )

                    # Создаем варианты ответов
                    for answer_data in question_data['answers']:
                        Answer.objects.create(
                            question=question,
                            answer_text=answer_data['text'],
                            is_correct=answer_data['is_correct'],
                            order=answer_data['order']
                        )

                    imported_count += 1

                    if self.verbose and idx % 10 == 0:
                        self.stdout.write(self.style.SUCCESS(f'[OK] Импортировано {idx} вопросов...'))

            except Exception as e:
                error_msg = f'Ошибка при импорте вопроса #{idx}: {e}'
                errors.append(error_msg)
                self.stdout.write(self.style.ERROR(f'[ERROR] {error_msg}'))

        # Итоги
        self.stdout.write(self.style.SUCCESS(f'\n{"="*70}'))
        if dry_run:
            self.stdout.write(self.style.WARNING('ТЕСТОВЫЙ РЕЖИМ - данные не сохранены'))
            self.stdout.write(f'Всего вопросов для импорта: {len(questions)}')
        else:
            self.stdout.write(self.style.SUCCESS(f'[OK] Успешно импортировано вопросов: {imported_count}'))

        if warnings:
            self.stdout.write(self.style.WARNING(f'\n[!] Предупреждений: {len(warnings)}'))
            for warning in warnings[:5]:  # Показываем первые 5
                self.stdout.write(self.style.WARNING(f'  - {warning}'))

        if errors:
            self.stdout.write(self.style.ERROR(f'\n[ERROR] Ошибок: {len(errors)}'))
            for error in errors[:5]:
                self.stdout.write(self.style.ERROR(f'  - {error}'))

        self.stdout.write(self.style.SUCCESS(f'{"="*70}\n'))

    def parse_questions_v2(self, doc):
        """
        Улучшенный парсер с поддержкой:
        - Автоопределения правильного ответа по жирному шрифту
        - Извлечения источника (пояснения)

        Ожидаемая структура:
        - Вопрос № X из Y (жирный)
        - Вопрос: текст вопроса
        - Варианты ответа:
        - Вариант 1 (обычный или ЖИРНЫЙ = правильный)
        - Вариант 2
        - ...
        - Источник: текст источника (пояснение к вопросу)
        """
        questions = []
        current_question = None
        current_state = None  # 'header', 'question', 'answers', 'source'

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if not text:
                continue

            # Проверяем жирность параграфа
            is_bold_para = self._is_paragraph_bold(para)

            if self.verbose and i < 50:
                bold_mark = '[B]' if is_bold_para else '   '
                self.stdout.write(f'{bold_mark} [{i}] {text[:60]}')

            # 1. Начало нового вопроса: "Вопрос № X из Y" или "Вопрос X из Y"
            if re.match(r'Вопрос\s*№?\s*\d+\s+из\s+\d+', text):
                # Сохраняем предыдущий вопрос
                if current_question and current_question['answers']:
                    # Форматируем текст вопроса перед сохранением
                    current_question['text'] = self._format_question_text(current_question['text'])
                    questions.append(current_question)

                current_question = {
                    'text': '',
                    'answers': [],
                    'explanation': '',
                    'answer_order': 1
                }
                current_state = 'header'
                continue

            # 2. Текст вопроса: "Вопрос: ..."
            if text.startswith('Вопрос:') and current_question is not None:
                # Убираем "Вопрос:" из начала
                question_text = text.replace('Вопрос:', '').strip()
                # Начинаем накапливать текст вопроса
                current_question['text'] = question_text
                current_state = 'question'
                continue

            # 2a. Продолжение текста вопроса (до "Варианты ответа:")
            # Накапливаем все строки между "Вопрос:" и "Варианты ответа:"
            if current_state == 'question' and current_question is not None:
                # Добавляем текст к вопросу через перевод строки, чтобы варианты А, Б, В, Г
                # отображались с новой строки
                if text and text != 'Варианты ответа:':
                    current_question['text'] += '\n' + text
                    continue

            # 3. Начало вариантов ответа
            if text == 'Варианты ответа:' and current_question is not None:
                current_state = 'answers'
                continue

            # 4. Источник (пояснение)
            if text.startswith('Источник:') and current_question is not None:
                source_text = text.replace('Источник:', '').strip()
                current_question['explanation'] = source_text
                current_state = 'source'
                continue

            # 5. Варианты ответов
            if current_state == 'answers' and current_question is not None:
                # Проверяем, жирный ли текст (правильный ответ)
                is_correct = self._is_paragraph_bold(para)

                # Убираем нумерацию вида "1. ", "2. " и т.д.
                answer_text = re.sub(r'^\d+\.\s*', '', text)
                # Форматируем текст ответа (убираем лишние символы ·)
                answer_text = self._format_answer_text(answer_text)

                current_question['answers'].append({
                    'text': answer_text,
                    'is_correct': is_correct,
                    'is_bold': is_correct,
                    'order': current_question['answer_order']
                })
                current_question['answer_order'] += 1

        # Добавляем последний вопрос
        if current_question and current_question['answers']:
            # Форматируем текст вопроса перед сохранением
            current_question['text'] = self._format_question_text(current_question['text'])
            questions.append(current_question)

        return questions

    def _is_paragraph_bold(self, para):
        """
        Проверяет, является ли весь параграф или его основная часть жирным.
        """
        if not para.runs:
            return False

        # Считаем длину жирного и обычного текста
        bold_length = 0
        total_length = 0

        for run in para.runs:
            text_len = len(run.text)
            total_length += text_len
            if run.bold:
                bold_length += text_len

        # Если больше 70% текста жирное, считаем параграф жирным
        if total_length > 0:
            return (bold_length / total_length) > 0.7

        return False

    def _format_question_text(self, text):
        """
        Форматирует текст вопроса для правильного отображения с переводами строк.
        Заменяет символы · перед А, Б, В, Г, Д и т.д. на переводы строк.

        Примеры:
        "С какого дня начинается действие·А – со дня начала" -> "С какого дня начинается действие\nА – со дня начала"
        """
        import re

        # Паттерн для замены: символ · перед заглавными русскими буквами А-Я
        # Заменяем "·А" на "\nА", "·Б" на "\nБ" и т.д.
        text = re.sub(r'·([А-ЯЁ][\s–—-])', r'\n\1', text)

        # Также обрабатываем случаи с цифрами: "·1.", "·2." и т.д.
        text = re.sub(r'·(\d+\.)', r'\n\1', text)

        # Убираем оставшиеся одиночные символы · в начале или конце
        text = text.strip('·').strip()

        return text

    def _format_answer_text(self, text):
        """
        Форматирует текст варианта ответа для правильного отображения.
        Убирает лишние символы · и форматирует текст.

        Примеры:
        "А·–·со·дня·начала·работы" -> "А – со дня начала работы"
        """
        # Убираем символы · в начале и конце
        text = text.strip('·').strip()

        # Заменяем множественные пробелы на одиночные
        text = re.sub(r'\s+', ' ', text)

        return text
